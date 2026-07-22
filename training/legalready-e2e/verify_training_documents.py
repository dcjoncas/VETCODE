#!/usr/bin/env python3
"""Structural QA for the generated LegalReady DOCX training package."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
EXPECTED = {
    "LegalReady-End-to-End-Training-Guide.docx": [
        "LegalReady End-to-End Training",
        "The eight-step exercise",
        "Automated validation result",
        "Match score",
        "63",
        "training-no-send",
    ],
    "Sample-JD-Legal-Operations-eDiscovery-Analyst.docx": [
        "Legal Operations & eDiscovery Analyst",
        "Python",
        "SQL",
        "Power BI",
        "SharePoint",
        "Adobe Acrobat",
    ],
    "Sample-Resume-Jordan-Ellis-Mitch.docx": ["Jordan Ellis", "mitch.blake@legalready.io", "Python", "Power BI"],
    "Sample-Resume-Jordan-Ellis-Michael.docx": ["Jordan Ellis", "michael.shrader@legalready.io", "Python", "Power BI"],
    "Sample-Resume-Jordan-Ellis-Kacey-Jo.docx": ["Jordan Ellis", "kacey-jo.hyde@legalready.io", "Python", "Power BI"],
}


def attr_int(node, name):
    value = node.get(qn(name)) if node is not None else None
    return int(value) if value is not None else None


def document_text(doc: Document) -> str:
    parts = [paragraph.text for paragraph in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)


def verify_table(table, file_name: str, table_index: int, issues: list[str]):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    layout = tbl_pr.find(qn("w:tblLayout"))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    grid_widths = [attr_int(col, "w:w") for col in table._tbl.tblGrid.findall(qn("w:gridCol"))]
    width = attr_int(tbl_w, "w:w")
    if attr_int(tbl_w, "w:w") is None or tbl_w.get(qn("w:type")) != "dxa":
        issues.append(f"{file_name} table {table_index}: missing DXA table width")
    if layout is None or layout.get(qn("w:type")) != "fixed":
        issues.append(f"{file_name} table {table_index}: table layout is not fixed")
    if tbl_ind is None or tbl_ind.get(qn("w:type")) != "dxa":
        issues.append(f"{file_name} table {table_index}: missing DXA table indent")
    if not grid_widths or any(value is None for value in grid_widths):
        issues.append(f"{file_name} table {table_index}: incomplete table grid")
    elif width != sum(grid_widths):
        issues.append(f"{file_name} table {table_index}: tblW {width} does not equal grid {sum(grid_widths)}")
    for row_index, row in enumerate(table._tbl.tr_lst):
        if len(row.tc_lst) != len(grid_widths):
            continue  # merged header/callout row
        for cell_index, tc in enumerate(row.tc_lst):
            tc_w = tc.tcPr.find(qn("w:tcW")) if tc.tcPr is not None else None
            if tc_w is None or tc_w.get(qn("w:type")) != "dxa":
                issues.append(f"{file_name} table {table_index} row {row_index} cell {cell_index}: missing DXA cell width")
            elif attr_int(tc_w, "w:w") != grid_widths[cell_index]:
                issues.append(
                    f"{file_name} table {table_index} row {row_index} cell {cell_index}: tcW {attr_int(tc_w, 'w:w')} != grid {grid_widths[cell_index]}"
                )


def verify_document(path: Path) -> dict:
    doc = Document(path)
    issues: list[str] = []
    text = document_text(doc)
    for needle in EXPECTED[path.name]:
        if needle not in text:
            issues.append(f"Missing expected text: {needle}")
    if "[YOUR" in text or "TBD" in text or "INSERT_EMAIL" in text:
        issues.append("Placeholder text remains")
    if any(paragraph.text.startswith(("- ", "• ")) for paragraph in doc.paragraphs):
        issues.append("Fake bullet paragraph detected")

    for index, section in enumerate(doc.sections):
        if round(section.page_width.inches, 2) != 8.5 or round(section.page_height.inches, 2) != 11.0:
            issues.append(f"Section {index}: page is not US Letter portrait")
        margins = [section.top_margin.inches, section.right_margin.inches, section.bottom_margin.inches, section.left_margin.inches]
        if any(round(value, 2) != 1.0 for value in margins):
            issues.append(f"Section {index}: margins are not one inch")
        if round(section.header_distance.inches, 3) != 0.492 or round(section.footer_distance.inches, 3) != 0.492:
            issues.append(f"Section {index}: header/footer distance mismatch")

    normal = doc.styles["Normal"]
    if normal.font.name != "Calibri" or round(normal.font.size.pt, 1) != 11.0:
        issues.append("Normal style does not use Calibri 11")
    expected_sizes = {"Heading 1": 16.0, "Heading 2": 13.0, "Heading 3": 12.0}
    for style_name, size in expected_sizes.items():
        if round(doc.styles[style_name].font.size.pt, 1) != size:
            issues.append(f"{style_name} size mismatch")

    for index, table in enumerate(doc.tables):
        verify_table(table, path.name, index, issues)

    return {
        "file": path.name,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "sections": len(doc.sections),
        "issues": issues,
        "ok": not issues,
    }


def main() -> int:
    results = []
    for name in EXPECTED:
        path = ROOT / name
        if not path.exists():
            results.append({"file": name, "ok": False, "issues": ["File missing"]})
            continue
        results.append(verify_document(path))
    report = {"ok": all(item["ok"] for item in results), "documents": results}
    output = ROOT / "qa" / "structural-qa.json"
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text(json.dumps(report, indent=2), encoding="utf-8")
    print(json.dumps(report, indent=2))
    return 0 if report["ok"] else 1


if __name__ == "__main__":
    raise SystemExit(main())
