#!/usr/bin/env python3
"""Build the LegalReady end-to-end training guide and reusable sample files."""

from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
RESULTS = ROOT / "results" / "legalready-e2e-latest.json"

NAVY = RGBColor(18, 58, 43)
GREEN = RGBColor(63, 126, 75)
LIGHT_GREEN = "EDF5EE"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = RGBColor(84, 97, 89)
BLACK = RGBColor(26, 31, 28)
WHITE = RGBColor(255, 255, 255)
RED = RGBColor(155, 28, 28)


TRAINEES = [
    ("Mitch Blake", "mitch.blake@legalready.io", "Sample-Resume-Jordan-Ellis-Mitch.docx"),
    ("Michael Shrader", "michael.shrader@legalready.io", "Sample-Resume-Jordan-Ellis-Michael.docx"),
    ("Kacey-Jo Hyde", "kacey-jo.hyde@legalready.io", "Sample-Resume-Jordan-Ellis-Kacey-Jo.docx"),
]


def set_run_font(run, name="Calibri", size=11, color=BLACK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_styles(doc: Document, preset: str):
    if preset == "compact_reference_guide":
        body_after, body_line = 6, 1.25
        h1 = (16, 18, 10, NAVY)
        h2 = (13, 14, 7, NAVY)
        h3 = (12, 10, 5, RGBColor(31, 77, 120))
    else:
        body_after, body_line = 6, 1.10
        h1 = (16, 16, 8, NAVY)
        h2 = (13, 12, 6, NAVY)
        h3 = (12, 8, 4, RGBColor(31, 77, 120))

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(body_after)
    normal.paragraph_format.line_spacing = body_line

    for style_name, values in zip(("Heading 1", "Heading 2", "Heading 3"), (h1, h2, h3)):
        style = doc.styles[style_name]
        size, before, after, color = values
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_name in ("List Bullet", "List Number"):
        style = doc.styles[list_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375 if preset == "compact_reference_guide" else 0.5)
        style.paragraph_format.first_line_indent = Inches(-0.188 if preset == "compact_reference_guide" else -0.25)
        style.paragraph_format.space_after = Pt(4 if preset == "compact_reference_guide" else 8)
        style.paragraph_format.line_spacing = 1.25 if preset == "compact_reference_guide" else 1.167


def configure_page(doc: Document):
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11)
        section.top_margin = Inches(1)
        section.right_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.header_distance = Inches(0.492)
        section.footer_distance = Inches(0.492)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_in, indent_dxa=120):
    widths_dxa = [round(width * 1440) for width in widths_in]
    total = sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            cell.width = Inches(widths_in[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    set_run_font(run, size=9, color=MID_GRAY)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def add_header_footer(doc: Document, left: str, right: str):
    for section in doc.sections:
        header = section.header
        p = header.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(left)
        set_run_font(run, size=9, color=MID_GRAY, bold=True)
        run = p.add_run(f"    {right}")
        set_run_font(run, size=9, color=MID_GRAY)
        add_page_number(section.footer.paragraphs[0])


def add_title_block(doc: Document, kicker: str, title: str, subtitle: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(kicker.upper())
    set_run_font(run, size=10, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(title)
    set_run_font(run, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run(subtitle)
    set_run_font(run, size=13, color=MID_GRAY)


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        set_run_font(run, bold=True)
        run = p.add_run(text[len(bold_prefix):])
        set_run_font(run)
    else:
        set_run_font(p.add_run(text))
    return p


def add_number(doc: Document, text: str):
    p = doc.add_paragraph(style="List Number")
    set_run_font(p.add_run(text))
    return p


def add_callout(doc: Document, label: str, text: str, fill=LIGHT_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [6.5], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label.upper())
    set_run_font(run, size=10, color=accent, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(text), size=10.5, color=BLACK)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def add_label_value_table(doc: Document, rows, widths=(1.875, 4.625), header=None):
    count = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=count, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, list(widths), indent_dxa=120)
    offset = 0
    if header:
        table.cell(0, 0).merge(table.cell(0, 1))
        set_cell_shading(table.cell(0, 0), LIGHT_GREEN)
        p = table.cell(0, 0).paragraphs[0]
        set_run_font(p.add_run(header), size=10.5, color=NAVY, bold=True)
        offset = 1
    for row_index, (label, value) in enumerate(rows, start=offset):
        set_cell_shading(table.cell(row_index, 0), LIGHT_GRAY)
        set_run_font(table.cell(row_index, 0).paragraphs[0].add_run(label), size=10, color=NAVY, bold=True)
        set_run_font(table.cell(row_index, 1).paragraphs[0].add_run(str(value)), size=10, color=BLACK)
    return table


def add_steps_table(doc: Document, rows):
    table = doc.add_table(rows=len(rows) + 1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [0.55, 1.35, 3.45, 1.15], indent_dxa=120)
    headers = ("#", "Screen", "Trainee action", "Pass evidence")
    for index, value in enumerate(headers):
        set_cell_shading(table.cell(0, index), "DDEADF")
        p = table.cell(0, index).paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if index in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT
        set_run_font(p.add_run(value), size=9.5, color=NAVY, bold=True)
    set_repeat_header(table.rows[0])
    for row_index, row in enumerate(rows, start=1):
        for col_index, value in enumerate(row):
            p = table.cell(row_index, col_index).paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_index in {0, 3} else WD_ALIGN_PARAGRAPH.LEFT
            set_run_font(p.add_run(str(value)), size=9.2, color=BLACK, bold=(col_index == 1))
    return table


def make_guide(path: Path):
    doc = Document()
    configure_styles(doc, "compact_reference_guide")
    configure_page(doc)
    add_header_footer(doc, "LEGALREADY TRAINING", "Law workflow | Internal practice")
    add_title_block(
        doc,
        "Workshop agenda",
        "LegalReady End-to-End Training",
        "A simple, repeatable exercise: JD to candidate, vetting, client package, scheduling, and status proof.",
    )

    metric = doc.add_table(rows=1, cols=4)
    set_table_geometry(metric, [1.625, 1.625, 1.625, 1.625], indent_dxa=120)
    for index, (top, bottom) in enumerate((("45-60 min", "Training time"), ("1", "Sample role"), ("1", "Sample candidate"), ("0", "Outside recipients"))):
        cell = metric.cell(0, index)
        set_cell_shading(cell, "E9F3EA")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_run_font(p.add_run(top), size=12, color=NAVY, bold=True)
        p = cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(bottom), size=8.5, color=MID_GRAY)

    doc.add_heading("Training outcome", level=1)
    add_callout(
        doc,
        "Definition of done",
        "The trainee creates the sample JD, uploads the matching résumé, finds and vets the candidate, prepares client communication, generates both interview drafts, records the schedule handoff, and proves the records in Status. Candidate, client, and interviewer addresses all use the trainee's own LegalReady email.",
    )
    add_bullet(doc, "Use only the supplied fictional candidate and fictional client.")
    add_bullet(doc, "Keep Domain set to Law on every screen.")
    add_bullet(doc, "Use the trainee's own LegalReady address for candidate, client, interviewer, and any additional attendee.")
    add_bullet(doc, "During training, generate drafts and archive the schedule. Do not send a live calendar invite unless the trainer explicitly authorizes it.")

    doc.add_heading("Trainee files and safe routing", level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    set_table_geometry(table, [1.35, 2.95, 2.20], indent_dxa=120)
    for col, value in enumerate(("Trainee", "Résumé file", "Use this address everywhere")):
        set_cell_shading(table.cell(0, col), "DDEADF")
        set_run_font(table.cell(0, col).paragraphs[0].add_run(value), size=9.5, color=NAVY, bold=True)
    set_repeat_header(table.rows[0])
    for row_index, (name, email, resume_file) in enumerate(TRAINEES, start=1):
        set_run_font(table.cell(row_index, 0).paragraphs[0].add_run(name), size=9.2, bold=True)
        set_run_font(table.cell(row_index, 1).paragraphs[0].add_run(resume_file), size=8.8)
        set_run_font(table.cell(row_index, 2).paragraphs[0].add_run(email), size=8.8)

    doc.add_page_break()
    doc.add_heading("The eight-step exercise", level=1)
    add_steps_table(
        doc,
        [
            ("1", "Job Descriptions", "Upload Sample-JD-Legal-Operations-eDiscovery-Analyst.docx. Client: Summit & Vale LLP - LegalReady Training. Title: Legal Operations & eDiscovery Analyst - Training.", "Saved JD ID"),
            ("2", "Find Candidate", "Upload the résumé assigned to you. The file already contains your LegalReady email. Wait for Profile Preview to open.", "Profile ID"),
            ("3", "Profile Preview", "Confirm Jordan Ellis, your LegalReady email, litigation experience, tools, and skills. Add to shortlist.", "Profile complete"),
            ("4", "Match to Role", "Load the saved training JD, rank profiles, find Jordan Ellis, record the score and top matches, then keep Jordan shortlisted.", "Candidate ranked"),
            ("5", "Client Communication", "Build the shortlist package. Client display name: Avery Stone (Training Client). Client email: your own LegalReady address.", "Draft package"),
            ("6A", "Candidate Review", "Candidate email and interviewer email must both be your own LegalReady address. Generate the email draft and archive the training schedule.", "Draft + archive"),
            ("6B", "Client Interview", "Client contact and candidate email must both be your own LegalReady address. Generate the client draft and archive the training schedule.", "Draft + archive"),
            ("7", "Status", "Open Status and verify both records, role, candidate, client, dates, message, and next action.", "Two records visible"),
        ],
    )

    doc.add_heading("Step-by-step trainer script", level=1)
    for text in [
        "Start at /ui/pages/job-descriptions.html?domain=law. Say: We begin with the client's need, not with a candidate.",
        "Save the sample role. Read back the client, title, core work, required skills, and success measures.",
        "Open Find Candidate and upload the assigned résumé. Explain that LegalReady extracts a structured profile from the file.",
        "On Profile Preview, confirm contact data, work history, skills, and evidence before matching.",
        "Open Match to Role, load the saved JD, and rank profiles. Ask the trainee to explain why Jordan Ellis fits and what still needs human verification.",
        "Shortlist the candidate and open Client Communication. Keep the fictional client display name but route it to the trainee's own email.",
        "Open Candidate Review. Generate the draft, review it, and save the schedule handoff. Do not send a real invite during the exercise.",
        "Switch to Client Interview. Generate the client draft with the same safe address, then save the second schedule handoff.",
        "Open Status. The trainee passes only after they can show both archived records and explain the next action.",
    ]:
        add_number(doc, text)

    doc.add_page_break()
    doc.add_heading("Vetting checklist", level=1)
    add_callout(doc, "Human decision required", "AI can extract, organize, rank, and draft. The trainee must verify facts, judge fit, correct errors, and approve any communication or scheduling action.", fill="FFF7E6", accent=RGBColor(122, 90, 0))
    checks = [
        "Identity and contact: name and LegalReady training email are correct.",
        "Relevant experience: at least three years of legal operations, litigation support, or eDiscovery work is visible.",
        "Data workflow: Python and SQL examples are supported by résumé evidence.",
        "Reporting: Power BI and Microsoft Excel experience is supported by specific results.",
        "Tools: SharePoint, Microsoft 365, and Adobe Acrobat are present.",
        "Confidentiality and client service are demonstrated.",
        "Gaps are recorded instead of invented. Suggested checks: jurisdiction-specific filing rules, compensation, start date, and reference verification.",
        "The trainee can explain the match score in plain language and does not treat the score as the final hiring decision.",
    ]
    for item in checks:
        add_bullet(doc, f"[ ] {item}")

    doc.add_heading("Pass/fail evidence", level=1)
    add_label_value_table(
        doc,
        [
            ("JD", "Training role is saved under the Law domain and can be reopened."),
            ("Candidate", "Jordan Ellis profile opens with the trainee's LegalReady email."),
            ("Match", "Candidate appears in ranked results and the trainee records score plus top matching evidence."),
            ("Shortlist", "Candidate is present in the client communication package."),
            ("Candidate review", "Draft is generated and a training-no-send archive record is visible."),
            ("Client interview", "Draft is generated and a training-no-send archive record is visible."),
            ("Routing", "Candidate, client, interviewer, and attendee addresses all equal the trainee's LegalReady email."),
            ("Safety", "No message or invite went to an outside address."),
        ],
        header="A trainee passes when every row is demonstrated",
    )

    doc.add_heading("Exact application path", level=1)
    for item in [
        "Job Descriptions: /ui/pages/job-descriptions.html?domain=law",
        "Find Candidate: /ui/pages/find-candidate.html?domain=law",
        "Match to Role: /ui/pages/match-role.html?domain=law",
        "Profile Preview: /ui/pages/profile-preview.html?domain=law",
        "Client Communication: /ui/pages/client-comm.html?domain=law",
        "Candidate Review: /ui/pages/schedule-interview.html?domain=law&interview=ready",
        "Client Interview: /ui/pages/schedule-interview.html?domain=law&interview=client",
        "Status: /ui/pages/status-tracker.html?domain=law",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()
    doc.add_heading("Automated validation result", level=1)
    if RESULTS.exists():
        result = json.loads(RESULTS.read_text(encoding="utf-8"))
        artifacts = result.get("artifacts") or {}
        match = artifacts.get("match") or {}
        archive_ids = artifacts.get("archive_ids") or []
        add_callout(
            doc,
            "Latest tested result",
            f"{'PASS' if result.get('ok') else 'FAIL'} - {result.get('run_id')}. The run used {result.get('email')} for candidate, client, interviewer, and attendee routing. It generated drafts and archive records but did not send email or create a calendar invite.",
            fill=LIGHT_GREEN if result.get("ok") else "FDECEC",
            accent=GREEN if result.get("ok") else RED,
        )
        add_label_value_table(
            doc,
            [
                ("Base URL", result.get("base_url", "")),
                ("Domain", result.get("domain", "")),
                ("JD ID", (artifacts.get("jd") or {}).get("jd_id", "")),
                ("Profile ID", (artifacts.get("profile") or {}).get("profile_id", "")),
                ("Match score", match.get("score", "Not returned")),
                ("Top matches", ", ".join((match.get("top_matches") or [])[:8]) or "Not returned"),
                ("Archive IDs", "\n".join(archive_ids) or "None"),
                ("No-send control", result.get("no_send_guarantee", "")),
            ],
            header="Latest proof packet",
        )
        doc.add_heading("Validation checks", level=2)
        for step in result.get("steps", []):
            add_bullet(doc, f"{'PASS' if step.get('ok') else 'FAIL'} - {step.get('name')}: {step.get('detail')}")
    else:
        add_callout(doc, "Validation pending", "Run run_legalready_e2e.py --execute after the sample files are generated. Rebuild this guide to embed the latest proof packet.", fill="FFF7E6", accent=RGBColor(122, 90, 0))

    doc.add_heading("Trainer closeout", level=1)
    add_bullet(doc, "Ask the trainee to explain the full flow without reading the guide.")
    add_bullet(doc, "Ask what the AI did and what still required human judgment.")
    add_bullet(doc, "Confirm the two Status records and same-email routing before marking complete.")
    add_bullet(doc, "Remove or clearly label training records if they should not remain in the operational list.")

    doc.core_properties.title = "LegalReady End-to-End Training Guide"
    doc.core_properties.subject = "Law-domain recruiting workflow training"
    doc.core_properties.author = "LegalReady Training"
    doc.core_properties.keywords = "LegalReady, training, recruiting, legal, candidate, scheduling"
    doc.save(path)


def make_jd(path: Path):
    doc = Document()
    configure_styles(doc, "standard_business_brief")
    configure_page(doc)
    add_header_footer(doc, "SAMPLE JOB DESCRIPTION", "Fictional training data")
    add_title_block(doc, "LegalReady training client", "Legal Operations & eDiscovery Analyst", "Summit & Vale LLP - fictional training role")
    add_label_value_table(
        doc,
        [
            ("Location", "Denver, Colorado - hybrid"),
            ("Employment", "Full-time"),
            ("Experience", "3+ years in legal operations or eDiscovery"),
            ("Reports to", "Director of Legal Operations"),
            ("Training note", "Fictional role. Use only in the Law training workflow."),
        ],
    )
    doc.add_heading("Role summary", level=1)
    doc.add_paragraph("Summit & Vale LLP is seeking a Legal Operations & eDiscovery Analyst to organize matter data, support discovery workflows, improve reporting, and help attorneys make reliable decisions from structured information. The role combines legal process knowledge with practical data and collaboration tools.")
    doc.add_heading("Core responsibilities", level=1)
    for item in [
        "Collect, normalize, and quality-check matter, discovery, and document-review data using Microsoft Excel, SQL, and Python.",
        "Build clear Power BI reports for matter status, discovery volumes, deadlines, review progress, and outside-counsel spend.",
        "Maintain secure legal workspaces in SharePoint and Microsoft 365 using consistent naming, permissions, version, and confidentiality controls.",
        "Prepare eDiscovery data sets, document productions, privilege logs, exhibits, and review-quality reports.",
        "Use Adobe Acrobat to prepare, redact, Bates-label, and quality-check production and filing documents.",
        "Translate attorney and client requests into documented workflows, data checks, and repeatable operating procedures.",
        "Communicate risks, exceptions, and next actions clearly to attorneys, clients, vendors, and the legal operations team.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Required qualifications", level=1)
    for item in [
        "Three or more years of legal operations, litigation support, eDiscovery, paralegal, or legal data experience.",
        "Hands-on Python and SQL experience for data cleanup, validation, or reporting.",
        "Strong Power BI and Microsoft Excel skills, including repeatable reporting and quality checks.",
        "Working knowledge of SharePoint, Microsoft 365, Adobe Acrobat, and secure document handling.",
        "Clear writing, careful quality control, professional client service, and sound judgment with confidential information.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Preferred qualifications", level=1)
    for item in [
        "Experience with Relativity, Microsoft Purview eDiscovery, Clio, or NetDocuments.",
        "Experience supporting commercial, employment, or regulatory matters.",
        "Ability to explain technical findings to lawyers and clients in plain language.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("First-90-day success measures", level=1)
    for item in [
        "Matter and discovery reports reconcile to source data with documented checks.",
        "Dashboards show accurate status, owners, risks, and next actions.",
        "Legal workspaces are complete, searchable, permissioned, and consistently organized.",
        "Attorneys and clients receive clear, timely findings and status updates.",
    ]:
        add_number(doc, item)
    doc.add_heading("Interview focus", level=1)
    doc.add_paragraph("Ask for specific examples of cleaning a difficult data set, writing a SQL or Python check, building a Power BI report, supporting an eDiscovery production, handling confidential material, and communicating a data risk to an attorney or client.")
    doc.core_properties.title = "Sample JD - Legal Operations & eDiscovery Analyst"
    doc.core_properties.author = "LegalReady Training"
    doc.save(path)


def make_resume(path: Path, trainee_name: str, email: str):
    doc = Document()
    configure_styles(doc, "standard_business_brief")
    configure_page(doc)
    add_header_footer(doc, "SAMPLE RESUME", f"Routes to {trainee_name}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    set_run_font(p.add_run("Jordan Ellis"), size=24, color=NAVY, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run("Legal Operations & eDiscovery Analyst"), size=13, color=GREEN, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(14)
    set_run_font(p.add_run(f"Denver, Colorado | {email} | (303) 555-0148"), size=10.5, color=MID_GRAY)

    add_callout(doc, "Training identity", f"Fictional candidate. All candidate and client communications in this exercise must route to {trainee_name} at {email}.")
    doc.add_heading("Professional summary", level=1)
    doc.add_paragraph("Legal operations and eDiscovery analyst with six years of experience supporting commercial, employment, and regulatory matters. Uses Python, SQL, Power BI, Microsoft Excel, SharePoint, Microsoft 365, and Adobe Acrobat to clean data, improve reporting, support discovery, and protect confidential information. Known for accurate work, calm prioritization, and clear follow-through.")
    doc.add_heading("Core skills", level=1)
    for item in [
        "Python for data cleanup, validation, reconciliation, and repeatable legal operations checks",
        "SQL for matter, billing, discovery, and document-review data queries",
        "Power BI dashboards; Microsoft Excel analysis; quality-control reporting",
        "SharePoint; Microsoft 365; Microsoft Teams; secure matter workspaces",
        "Adobe Acrobat; eDiscovery; document production; redaction; Bates labeling",
        "Client service; confidentiality; workflow documentation; plain-language findings",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Professional experience", level=1)
    p = doc.add_paragraph()
    set_run_font(p.add_run("Legal Operations & eDiscovery Analyst | Front Range Legal Partners | Denver, CO"), size=11, color=NAVY, bold=True)
    set_run_font(p.add_run("  |  2021-Present"), size=10, color=MID_GRAY)
    for item in [
        "Support legal operations and eDiscovery work across commercial, employment, and regulatory matters.",
        "Use Python and SQL to clean, reconcile, and validate matter, billing, custodian, and document-review data.",
        "Build Power BI and Microsoft Excel reports that show review progress, production volumes, deadlines, owners, risks, and next actions.",
        "Maintain secure SharePoint and Microsoft 365 workspaces with consistent permissions, naming, version control, and confidentiality practices.",
        "Prepare eDiscovery productions, privilege logs, redactions, Bates labels, and quality checks using Adobe Acrobat and review-platform exports.",
        "Reduced weekly reporting time by 40% after automating reconciliation checks and publishing a reusable Power BI dashboard.",
    ]:
        add_bullet(doc, item)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    set_run_font(p.add_run("Litigation Support Coordinator | Holloway & Pierce LLP | Denver, CO"), size=11, color=NAVY, bold=True)
    set_run_font(p.add_run("  |  2018-2021"), size=10, color=MID_GRAY)
    for item in [
        "Supported litigation teams with document collections, discovery tracking, production logs, reporting, and matter-file maintenance.",
        "Used Microsoft Excel, SharePoint, Microsoft 365, and Adobe Acrobat to organize, quality-check, and deliver legal work product.",
        "Prepared discovery shells, subpoenas, records requests, service packages, deposition exhibits, and client status materials.",
        "Created a matter-closing and data-retention checklist that improved file completeness and reduced missing-document follow-up.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Education and professional development", level=1)
    doc.add_paragraph("Bachelor of Science, Information Systems - Metropolitan State University of Denver")
    doc.add_paragraph("Certificate, Paralegal Studies - Community College of Denver")
    doc.add_paragraph("Continuing education: eDiscovery fundamentals; ethical handling of client information; legal data quality")
    doc.add_heading("Tools", level=1)
    doc.add_paragraph("Python, SQL, Power BI, Microsoft Excel, SharePoint, Microsoft 365, Microsoft Teams, Adobe Acrobat, Relativity, Clio, NetDocuments")
    doc.add_heading("Training-only references", level=1)
    doc.add_paragraph("Available during the exercise. Do not contact real employers or references; all information in this résumé is fictional.")
    doc.core_properties.title = f"Sample Resume - Jordan Ellis - {trainee_name}"
    doc.core_properties.subject = "LegalReady training resume"
    doc.core_properties.author = "LegalReady Training"
    doc.save(path)


def main():
    ROOT.mkdir(parents=True, exist_ok=True)
    make_jd(ROOT / "Sample-JD-Legal-Operations-eDiscovery-Analyst.docx")
    for name, email, filename in TRAINEES:
        make_resume(ROOT / filename, name, email)
    make_guide(ROOT / "LegalReady-End-to-End-Training-Guide.docx")
    print("Generated LegalReady training guide, sample JD, and three same-email résumé files.")


if __name__ == "__main__":
    main()
