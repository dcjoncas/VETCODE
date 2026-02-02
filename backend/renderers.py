from docx import Document
from docx.shared import Pt
from html import escape
import html

def _h(s): 
    return escape(s or "")

def profile_to_html(p: dict) -> str:
    skills = p.get("skills", {})
    scores = p.get("scores", {})
    contact = p.get("contact", {})
    summary = p.get("summary", {})
    meta = p.get("meta", {})
    def chips(items):
        if not items: return "<span class='muted'>—</span>"
        return " ".join([f"<span class='chip'>{_h(x)}</span>" for x in items])
    html = f"""
<!DOCTYPE html>
<html><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>DevReady Profile - {_h(contact.get('full_name','Candidate'))}</title>
<style>
:root{{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;background:#f3f4f6;color:#0f172a}}
body{{margin:0;padding:24px}}
.card{{background:#fff;border-radius:16px;box-shadow:0 12px 30px rgba(2,6,23,.08);padding:18px;max-width:1000px;margin:0 auto}}
h1{{margin:0;font-size:26px}}
.meta{{color:#475569;font-size:13px;margin-top:6px}}
.section{{margin-top:16px}}
h2{{font-size:14px;margin:0 0 8px;color:#334155;text-transform:uppercase;letter-spacing:.06em}}
.row{{display:flex;gap:10px;flex-wrap:wrap}}
.kv{{background:#f8fafc;border:1px solid #e2e8f0;border-radius:12px;padding:10px 12px;font-size:13px}}
.k{{font-weight:800;color:#0f172a}}
.muted{{color:#64748b}}
.chip{{display:inline-block;background:#0b1220;color:#fff;border-radius:999px;padding:6px 10px;margin:4px 6px 0 0;font-size:12px}}
.scoregrid{{display:grid;grid-template-columns:repeat(3,minmax(180px,1fr));gap:10px}}
.score{{background:#0b1220;color:#fff;border-radius:14px;padding:12px}}
.score .lbl{{opacity:.75;font-size:12px;letter-spacing:.02em}}
.score .val{{font-size:26px;font-weight:900}}
.score .rat{{opacity:.85;font-size:12px;margin-top:6px}}
@media (max-width:900px){{.scoregrid{{grid-template-columns:1fr}}}}
</style></head>
<body>
  <div class="card">
    <h1>{_h(contact.get('full_name','Candidate'))}</h1>
    <div class="meta">Profile ID: {_h(meta.get('profile_id',''))} • Created: {_h(meta.get('created_at',''))}</div>

    <div class="section">
      <h2>Contact</h2>
      <div class="row">
        <div class="kv"><span class="k">Email:</span> {_h(contact.get('email',''))}</div>
        <div class="kv"><span class="k">Phone:</span> {_h(contact.get('phone',''))}</div>
        <div class="kv"><span class="k">Location:</span> {_h(contact.get('location',''))}</div>
        <div class="kv"><span class="k">LinkedIn:</span> {_h(contact.get('linkedin',''))}</div>
      </div>
    </div>

    <div class="section">
      <h2>Summary</h2>
      <div class="kv"><span class="k">Headline:</span> {_h(summary.get('headline',''))}</div>
      <div class="kv" style="margin-top:10px">{_h(summary.get('overview',''))}</div>
    </div>

    <div class="section">
      <h2>Scores</h2>
      <div class="scoregrid">
        <div class="score"><div class="lbl">OVERALL TECHNICAL</div><div class="val">{scores.get('overall_technical',{}).get('score','—')}/10</div><div class="rat">{_h(scores.get('overall_technical',{}).get('rationale',''))}</div></div>
        <div class="score"><div class="lbl">BACKEND</div><div class="val">{scores.get('backend',{}).get('score','—')}/10</div><div class="rat">{_h(scores.get('backend',{}).get('rationale',''))}</div></div>
        <div class="score"><div class="lbl">FRONTEND</div><div class="val">{scores.get('frontend',{}).get('score','—')}/10</div><div class="rat">{_h(scores.get('frontend',{}).get('rationale',''))}</div></div>
      </div>
    </div>

    <div class="section">
      <h2>Skills</h2>
      <div class="kv"><span class="k">Languages</span><div>{chips(skills.get('languages',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Backend</span><div>{chips(skills.get('backend',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Frontend</span><div>{chips(skills.get('frontend',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Cloud/DevOps</span><div>{chips(skills.get('cloud_devops',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Data</span><div>{chips(skills.get('data',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Testing</span><div>{chips(skills.get('testing',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><span class="k">Security</span><div>{chips(skills.get('security',[]))}</div></div>
    </div>
  </div>
</body></html>
"""
    return html.strip()

def profile_to_docx(p: dict, out_path: str):
    contact = p.get("contact", {})
    summary = p.get("summary", {})
    skills = p.get("skills", {})
    scores = p.get("scores", {})
    meta = p.get("meta", {})

    doc = Document()
    title = doc.add_paragraph(contact.get("full_name","Candidate"))
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.bold = True

    doc.add_paragraph(f"Profile ID: {meta.get('profile_id','')}")
    doc.add_paragraph(f"Created: {meta.get('created_at','')}")

    doc.add_heading("Contact", level=2)
    doc.add_paragraph(f"Email: {contact.get('email','')}")
    doc.add_paragraph(f"Phone: {contact.get('phone','')}")
    doc.add_paragraph(f"Location: {contact.get('location','')}")
    doc.add_paragraph(f"LinkedIn: {contact.get('linkedin','')}")

    doc.add_heading("Summary", level=2)
    doc.add_paragraph(f"Headline: {summary.get('headline','')}")
    doc.add_paragraph(summary.get("overview",""))

    doc.add_heading("Scores (0-10)", level=2)
    for k in ["overall_technical","backend","frontend","cloud_devops","data","testing"]:
        s = scores.get(k,{})
        doc.add_paragraph(f"{k.replace('_',' ').title()}: {s.get('score','')} — {s.get('rationale','')}")

    doc.add_heading("Skills", level=2)
    for bucket in ["languages","backend","frontend","cloud_devops","data","testing","security","other"]:
        vals = skills.get(bucket,[]) or []
        doc.add_paragraph(f"{bucket.replace('_',' ').title()}: " + (", ".join(vals) if vals else "—"))

    doc.save(out_path)

def jd_to_html(jd: dict) -> str:
    def chips(items):
        if not items: return "<span class='muted'>—</span>"
        return " ".join([f"<span class='chip'>{_h(x)}</span>" for x in items])
    skills = jd.get("jd_skills", {}) or {}
    html = f"""
<!DOCTYPE html>
<html><head><meta charset="utf-8"/><meta name="viewport" content="width=device-width,initial-scale=1"/>
<title>Job Description - {_h(jd.get('title',''))}</title>
<style>
:root{{font-family:system-ui,-apple-system,Segoe UI,Roboto,Arial,sans-serif;background:#f3f4f6;color:#0f172a}}
body{{margin:0;padding:24px}}
.card{{background:#fff;border-radius:16px;box-shadow:0 12px 30px rgba(2,6,23,.08);padding:18px;max-width:1000px;margin:0 auto}}
h1{{margin:0;font-size:22px}}
.meta{{color:#475569;font-size:13px;margin-top:6px}}
.section{{margin-top:16px}}
h2{{font-size:14px;margin:0 0 8px;color:#334155;text-transform:uppercase;letter-spacing:.06em}}
.kv{{background:#f8fafc;border:1px solid #e2e8f0;border-radius:12px;padding:10px 12px;font-size:13px;white-space:pre-wrap}}
.muted{{color:#64748b}}
.chip{{display:inline-block;background:#0b1220;color:#fff;border-radius:999px;padding:6px 10px;margin:4px 6px 0 0;font-size:12px}}
</style></head>
<body>
  <div class="card">
    <h1>{_h(jd.get('title',''))}</h1>
    <div class="meta">Company: {_h(jd.get('company',''))} • JD ID: {_h(jd.get('jd_id',''))} • Created: {_h(jd.get('created_at',''))}</div>
    <div class="section">
      <h2>Skill Buckets</h2>
      <div class="kv"><strong>Languages</strong><div>{chips(skills.get('languages',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><strong>Backend</strong><div>{chips(skills.get('backend',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><strong>Frontend</strong><div>{chips(skills.get('frontend',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><strong>Cloud/DevOps</strong><div>{chips(skills.get('cloud_devops',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><strong>Data</strong><div>{chips(skills.get('data',[]))}</div></div>
      <div class="kv" style="margin-top:10px"><strong>Testing</strong><div>{chips(skills.get('testing',[]))}</div></div>
    </div>
    <div class="section">
      <h2>Job Description</h2>
      <div class="kv">{_h(jd.get('jd_text',''))}</div>
    </div>
  </div>
</body></html>
"""
    return html.strip()

def jd_to_docx(jd: dict, out_path: str):
    doc = Document()
    title = doc.add_paragraph(jd.get("title","Job Description"))
    title.runs[0].font.size = Pt(18)
    title.runs[0].font.bold = True
    doc.add_paragraph(f"Company: {jd.get('company','')}")
    doc.add_paragraph(f"JD ID: {jd.get('jd_id','')}")
    doc.add_paragraph(f"Created: {jd.get('created_at','')}")

    doc.add_heading("Skill Buckets", level=2)
    skills = jd.get("jd_skills", {}) or {}
    for bucket in ["languages","backend","frontend","cloud_devops","data","testing","security"]:
        vals = skills.get(bucket,[]) or []
        doc.add_paragraph(f"{bucket.replace('_',' ').title()}: " + (", ".join(vals) if vals else "—"))

    doc.add_heading("Job Description", level=2)
    doc.add_paragraph(jd.get("jd_text",""))

    doc.save(out_path)


def match_report_to_html(profile: dict, jd: dict, scorecard: dict, interview: dict, explain: dict) -> str:
    """Build a shareable HTML recommendation pack (profile + match artifacts)."""
    # Defensive defaults
    p_name = profile.get("name") or "Candidate"
    p_email = profile.get("email") or ""
    p_loc = profile.get("location") or ""
    company = jd.get("company") or ""
    title = jd.get("title") or "Job Description"
    jd_id = jd.get("jd_id") or ""
    profile_id = profile.get("profile_id") or ""
    match_score = explain.get("match_score") or explain.get("matchScore") or 0

    s10 = (scorecard.get("scores_out_of_10") or scorecard.get("scores") or {})
    tech_obj = s10.get("technical", "—"); tech = (tech_obj.get("score") if isinstance(tech_obj, dict) else tech_obj) if tech_obj is not None else "—"
    func_obj = s10.get("functional", "—"); func = (func_obj.get("score") if isinstance(func_obj, dict) else func_obj) if func_obj is not None else "—"
    biz_obj = s10.get("business", "—"); biz = (biz_obj.get("score") if isinstance(biz_obj, dict) else biz_obj) if biz_obj is not None else "—"
    vertical = scorecard.get("vertical") or "—"

    top_matches = explain.get("top_matches") or []
    notable_gaps = explain.get("notable_gaps") or []
    pros = scorecard.get("pros") or []
    diffs = scorecard.get("differentiators") or []
    gaps = scorecard.get("gaps") or []
    cons = scorecard.get("cons") or []
    questions = interview.get("questions") or []

    excerpt = explain.get("client_excerpt") or ""
    draft_email = explain.get("draft_client_email") or ""

    def esc(x: str) -> str:
        # escape was imported from html; avoid NameError on missing `html` module
        return escape(str(x if x is not None else ""))

    def li(items):
        if not items:
            return "<li><span class='muted'>—</span></li>"
        return "\n".join([f"<li>{esc(i)}</li>" for i in items])

    # Basic styling matches the DevReady clean look
    return f"""<!doctype html>
<html lang='en'>
<head>
  <meta charset='utf-8'/>
  <meta name='viewport' content='width=device-width, initial-scale=1'/>
  <title>DevReady Recommendation — {esc(p_name)} — {esc(company)} {esc(title)}</title>
  <style>
    :root {{
      --dr-green:#7cc043; --bg:#f3f4f6; --panel:#ffffff; --text:#0f172a; --muted:#6b7280; --line:#e5e7eb;
      font-family: system-ui, -apple-system, Segoe UI, Roboto, sans-serif;
    }}
    body{{margin:0; background:var(--bg); color:var(--text); padding:24px}}
    .shell{{max-width:980px; margin:0 auto}}
    .card{{background:var(--panel); border:1px solid var(--line); border-radius:16px; padding:16px; box-shadow:0 10px 30px rgba(15,23,42,.06)}}
    .h1{{font-size:20px; font-weight:900; margin:0 0 6px}}
    .muted{{color:var(--muted)}}
    .grid{{display:grid; grid-template-columns:1fr; gap:14px}}
    @media (min-width:900px){{ .grid{{grid-template-columns:1fr 1fr}} }}
    .kpis{{display:flex; gap:10px; flex-wrap:wrap; margin-top:10px}}
    .kpi{{flex:1; min-width:170px; border:1px solid var(--line); border-radius:14px; padding:10px; background:#f9fafb}}
    .kpi .label{{font-size:12px; color:var(--muted)}}
    .kpi .val{{font-size:22px; font-weight:900}}
    .bar{{height:8px; background:#e5e7eb; border-radius:999px; overflow:hidden; margin-top:6px}}
    .bar > div{{height:100%; background:var(--dr-green); width:0}}
    .sectionTitle{{font-weight:900; margin:0 0 8px}}
    ul{{margin:6px 0 0 18px}}
    pre{{white-space:pre-wrap; background:#f9fafb; border:1px solid var(--line); border-radius:12px; padding:10px; font-family: ui-monospace, SFMono-Regular, Menlo, monospace; font-size:12px}}
    .pill{{display:inline-flex; align-items:center; padding:2px 10px; border-radius:999px; background:#e8f7e8; border:1px solid var(--line); font-size:12px; font-weight:700}}
    .row{{display:flex; justify-content:space-between; gap:10px; flex-wrap:wrap}}
    .small{{font-size:12px}}
  </style>
</head>
<body>
<div class='shell'>
  <div class='card'>
    <div class='row'>
      <div>
        <div class='h1'>Recommendation Pack</div>
        <div class='muted'>Profile {esc(profile_id)} • JD {esc(jd_id)}</div>
      </div>
      <div class='pill'>Technology domain</div>
    </div>
    <hr style='border:none;border-top:1px solid var(--line); margin:12px 0'/>
    <div class='row'>
      <div>
        <div style='font-size:18px;font-weight:900'>{esc(p_name)}</div>
        <div class='muted'>{esc(p_email)}{' • ' + esc(p_loc) if p_loc else ''}</div>
      </div>
      <div style='text-align:right'>
        <div class='small muted'>Role</div>
        <div style='font-weight:900'>{esc(company)} — {esc(title)}</div>
      </div>
    </div>

    <div class='kpis'>
      <div class='kpi'>
        <div class='label'>Match score (out of 100)</div>
        <div class='val'>{esc(round(float(match_score),1) if isinstance(match_score,(int,float)) else match_score)}</div>
        <div class='bar'><div style='width:{min(max(float(match_score),0),100)}%'></div></div>
      </div>
      <div class='kpi'>
        <div class='label'>Technical (out of 10)</div>
        <div class='val'>{esc(tech)}</div>
        <div class='bar'><div style='width:{min(max(float(tech) if str(tech).replace('.','',1).isdigit() else 0,0),10)/10*100}%'></div></div>
      </div>
      <div class='kpi'>
        <div class='label'>Functional (out of 10)</div>
        <div class='val'>{esc(func)}</div>
        <div class='bar'><div style='width:{min(max(float(func) if str(func).replace('.','',1).isdigit() else 0,0),10)/10*100}%'></div></div>
      </div>
      <div class='kpi'>
        <div class='label'>Business (out of 10)</div>
        <div class='val'>{esc(biz)}</div>
        <div class='bar'><div style='width:{min(max(float(biz) if str(biz).replace('.','',1).isdigit() else 0,0),10)/10*100}%'></div></div>
      </div>
    </div>

    <div style='margin-top:12px' class='grid'>
      <div class='card' style='box-shadow:none'>
        <div class='sectionTitle'>Top matches</div>
        <ul>{li(top_matches)}</ul>
        <div style='margin-top:10px' class='sectionTitle'>Notable gaps</div>
        <ul>{li(notable_gaps)}</ul>
      </div>
      <div class='card' style='box-shadow:none'>
        <div class='sectionTitle'>Scorecard</div>
        <div class='muted small'>Most proficient vertical: <b>{esc(vertical)}</b></div>
        <div style='margin-top:8px' class='row'>
          <div style='flex:1'>
            <div class='small muted'>Pros</div>
            <ul>{li(pros)}</ul>
            <div style='margin-top:10px' class='small muted'>Differentiators</div>
            <ul>{li(diffs)}</ul>
          </div>
          <div style='flex:1'>
            <div class='small muted'>Cons / Risks</div>
            <ul>{li(cons)}</ul>
            <div style='margin-top:10px' class='small muted'>Gaps to validate</div>
            <ul>{li(gaps)}</ul>
          </div>
        </div>
      </div>

      <div class='card' style='box-shadow:none'>
        <div class='sectionTitle'>Client-ready excerpt</div>
        <div>{esc(excerpt) or '<span class="muted">—</span>'}</div>
        <div style='margin-top:10px' class='sectionTitle'>Interview questions</div>
        <ol>{''.join([f'<li>{esc(q)}</li>' for q in questions]) or '<li class="muted">—</li>'}</ol>
      </div>

      <div class='card' style='box-shadow:none'>
        <div class='sectionTitle'>Draft client email</div>
        <pre>{esc(draft_email)}</pre>
      </div>
    </div>

    <div class='muted small' style='margin-top:10px'>Generated by DevReady Vetting (local).</div>
  </div>
</div>
</body>
</html>"""


def match_report_to_docx(out_path: str, profile: dict, jd: dict, scorecard: dict, interview: dict, explain: dict) -> str:
    """Write a DOCX recommendation pack and return file path."""
    doc = Document()

    p_name = profile.get("name") or "Candidate"
    p_email = profile.get("email") or ""
    p_loc = profile.get("location") or ""
    company = jd.get("company") or ""
    title = jd.get("title") or "Job Description"
    match_score = explain.get("match_score") or explain.get("matchScore") or 0

    doc.add_heading("DevReady Recommendation Pack", level=1)
    doc.add_paragraph(f"Candidate: {p_name}")
    if p_email:
        doc.add_paragraph(f"Email: {p_email}")
    if p_loc:
        doc.add_paragraph(f"Location: {p_loc}")
    doc.add_paragraph(f"Role: {company} — {title}")
    doc.add_paragraph(f"Match score (out of 100): {match_score}")

    s10 = (scorecard.get("scores_out_of_10") or scorecard.get("scores") or {})
    tech_obj = s10.get("technical", "—"); tech = (tech_obj.get("score") if isinstance(tech_obj, dict) else tech_obj) if tech_obj is not None else "—"
    func_obj = s10.get("functional", "—"); func = (func_obj.get("score") if isinstance(func_obj, dict) else func_obj) if func_obj is not None else "—"
    biz_obj = s10.get("business", "—"); biz = (biz_obj.get("score") if isinstance(biz_obj, dict) else biz_obj) if biz_obj is not None else "—"

    doc.add_heading("Scorecard (out of 10)", level=2)
    table = doc.add_table(rows=2, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Technical"
    hdr[1].text = "Functional"
    hdr[2].text = "Business"
    row = table.rows[1].cells
    row[0].text = str(tech)
    row[1].text = str(func)
    row[2].text = str(biz)

    vertical = scorecard.get("vertical") or "—"
    doc.add_paragraph(f"Most proficient vertical: {vertical}")

    def add_bullets(title_text: str, items):
        doc.add_heading(title_text, level=3)
        if not items:
            doc.add_paragraph("—")
            return
        for it in items:
            doc.add_paragraph(str(it), style="List Bullet")

    add_bullets("Top matches", explain.get("top_matches") or [])
    add_bullets("Notable gaps", explain.get("notable_gaps") or [])
    add_bullets("Pros", scorecard.get("pros") or [])
    add_bullets("Differentiators", scorecard.get("differentiators") or [])
    add_bullets("Cons / Risks", scorecard.get("cons") or [])
    add_bullets("Gaps to validate", scorecard.get("gaps") or [])

    doc.add_heading("Client-ready excerpt", level=2)
    doc.add_paragraph(explain.get("client_excerpt") or "—")

    doc.add_heading("Interview questions", level=2)
    qs = (interview.get("questions") or [])
    if qs:
        for q in qs:
            doc.add_paragraph(str(q), style="List Number")
    else:
        doc.add_paragraph("—")

    doc.add_heading("Draft client email", level=2)
    doc.add_paragraph(explain.get("draft_client_email") or "—")

    doc.save(out_path)
    return out_path